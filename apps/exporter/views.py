from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework import status
from .services import export_weekly_report_pdf, export_weekly_report_docx, export_general_report_pdf, export_general_report_docx
from apps.reports.models import WeeklyReport, GeneralReport, DailyReport
from apps.core.permissions import IsOwnerOrReadOnly
from django.http import HttpResponse


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def export_weekly_report_pdf_view(request, weekly_report_id):
    """Export weekly report as PDF"""
    try:
        weekly_report = WeeklyReport.objects.get(id=weekly_report_id, student=request.user)
    except WeeklyReport.DoesNotExist:
        return Response({
            'success': False,
            'message': 'Weekly report not found'
        }, status=404)
    
    try:
        pdf_content = export_weekly_report_pdf(weekly_report)
        response = HttpResponse(pdf_content, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="weekly_report_{weekly_report.week_number}.pdf"'
        return response
    except Exception as e:
        return Response({
            'success': False,
            'message': f'Error generating PDF: {str(e)}'
        }, status=500)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def export_weekly_report_docx_view(request, weekly_report_id):
    """Export weekly report as DOCX"""
    try:
        weekly_report = WeeklyReport.objects.get(id=weekly_report_id, student=request.user)
    except WeeklyReport.DoesNotExist:
        return Response({
            'success': False,
            'message': 'Weekly report not found'
        }, status=404)
    
    try:
        docx_content = export_weekly_report_docx(weekly_report)
        response = HttpResponse(docx_content, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="weekly_report_{weekly_report.week_number}.docx"'
        return response
    except Exception as e:
        return Response({
            'success': False,
            'message': f'Error generating DOCX: {str(e)}'
        }, status=500)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def export_daily_report_pdf(request, daily_report_id):
    """Export daily report as PDF"""
    try:
        daily_report = DailyReport.objects.get(id=daily_report_id, student=request.user)
    except DailyReport.DoesNotExist:
        return Response({
            'success': False,
            'message': 'Daily report not found'
        }, status=404)
    
    # For daily reports, we'll create a simple PDF
    from django.template.loader import render_to_string
    from weasyprint import HTML, CSS
    from io import BytesIO
    from django.http import HttpResponse
    from django.utils import timezone
    
    context = {
        'report': daily_report,
        'student': daily_report.student,
        'profile': daily_report.student.profile,
        'company': daily_report.student.profile.company,
        'export_date': timezone.now().strftime('%B %d, %Y')
    }
    
    html_string = render_to_string('exports/daily_report.html', context)
    
    css_string = '''
    @page {
        size: A4;
        margin: 1in;
    }
    body {
        font-family: 'Times New Roman', serif;
        font-size: 12pt;
        line-height: 1.5;
        color: #333;
    }
    .header {
        text-align: center;
        margin-bottom: 30px;
        border-bottom: 2px solid #333;
        padding-bottom: 20px;
    }
    .section {
        margin-bottom: 25px;
    }
    .section-title {
        font-size: 14pt;
        font-weight: bold;
        color: #2c3e50;
        border-bottom: 1px solid #bdc3c7;
        padding-bottom: 5px;
        margin-bottom: 15px;
    }
    '''
    
    html = HTML(string=html_string)
    css = CSS(string=css_string)
    
    pdf_buffer = BytesIO()
    html.write_pdf(pdf_buffer, stylesheets=[css])
    pdf_buffer.seek(0)
    
    response = HttpResponse(pdf_buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="daily_report_{daily_report.date}.pdf"'
    
    return response


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def export_daily_report_docx(request, daily_report_id):
    """Export daily report as DOCX"""
    try:
        daily_report = DailyReport.objects.get(id=daily_report_id, student=request.user)
    except DailyReport.DoesNotExist:
        return Response({
            'success': False,
            'message': 'Daily report not found'
        }, status=404)
    
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO
    from django.http import HttpResponse
    from django.utils import timezone
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Header
    header = doc.add_heading('DAILY TRAINING REPORT', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Student and company info
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ('Student Name:', daily_report.student.get_full_name()),
        ('Student ID:', daily_report.student.profile.student_id),
        ('Program:', daily_report.student.profile.get_program_display()),
        ('Company:', daily_report.student.profile.company.name),
        ('Date:', daily_report.date.strftime('%B %d, %Y')),
    ]
    
    for i, (label, value) in enumerate(info_data):
        info_table.cell(i, 0).text = label
        info_table.cell(i, 1).text = value
    
    doc.add_paragraph()
    
    # Daily Activities
    doc.add_heading('Daily Activities', level=1)
    doc.add_paragraph(daily_report.description)
    
    # Hours Spent
    doc.add_heading('Hours Spent', level=1)
    doc.add_paragraph(f"Total hours worked: {daily_report.hours_spent}")
    
    # Skills Learned
    if daily_report.skills_learned:
        doc.add_heading('Skills Learned', level=1)
        doc.add_paragraph(daily_report.skills_learned)
    
    # Challenges Faced
    if daily_report.challenges_faced:
        doc.add_heading('Challenges Faced', level=1)
        doc.add_paragraph(daily_report.challenges_faced)
    
    # Supervisor Feedback
    if daily_report.supervisor_feedback:
        doc.add_heading('Supervisor Feedback', level=1)
        doc.add_paragraph(daily_report.supervisor_feedback)
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph(f"Generated on {timezone.now().strftime('%B %d, %Y')}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to buffer
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    response = HttpResponse(doc_buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="daily_report_{daily_report.date}.docx"'
    
    return response


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def export_general_report_pdf_view(request):
    """Export general report as PDF"""
    try:
        general_report = GeneralReport.objects.get(user=request.user)
    except GeneralReport.DoesNotExist:
        return Response({
            'success': False,
            'message': 'General report not found'
        }, status=404)
    
    try:
        pdf_content = export_general_report_pdf(general_report)
        response = HttpResponse(pdf_content, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="general_report.pdf"'
        return response
    except Exception as e:
        return Response({
            'success': False,
            'message': f'Error generating PDF: {str(e)}'
        }, status=500)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def export_general_report_docx_view(request):
    """Export general report as DOCX"""
    try:
        general_report = GeneralReport.objects.get(user=request.user)
    except GeneralReport.DoesNotExist:
        return Response({
            'success': False,
            'message': 'General report not found'
        }, status=404)
    
    try:
        docx_content = export_general_report_docx(general_report)
        response = HttpResponse(docx_content, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="general_report.docx"'
        return response
    except Exception as e:
        return Response({
            'success': False,
            'message': f'Error generating DOCX: {str(e)}'
        }, status=500)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def bulk_export(request):
    """Bulk export multiple reports"""
    report_ids = request.data.get('report_ids', [])
    export_type = request.data.get('type', 'pdf')  # pdf or docx
    report_type = request.data.get('report_type', 'weekly')  # daily, weekly, general
    
    if not report_ids:
        return Response({
            'success': False,
            'message': 'Report IDs are required'
        }, status=400)
    
    # This is a simplified bulk export - in a real implementation,
    # you might want to create a zip file with multiple reports
    if report_type == 'weekly':
        try:
            weekly_report = WeeklyReport.objects.get(id=report_ids[0], student=request.user)
            if export_type == 'pdf':
                pdf_content = export_weekly_report_pdf(weekly_report)
                response = HttpResponse(pdf_content, content_type='application/pdf')
                response['Content-Disposition'] = f'attachment; filename="weekly_report_{weekly_report.week_number}.pdf"'
                return response
            else:
                docx_content = export_weekly_report_docx(weekly_report)
                response = HttpResponse(docx_content, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename="weekly_report_{weekly_report.week_number}.docx"'
                return response
        except WeeklyReport.DoesNotExist:
            return Response({
                'success': False,
                'message': 'Weekly report not found'
            }, status=404)
    
    return Response({
        'success': False,
        'message': 'Bulk export not implemented for this report type'
    }, status=400) 