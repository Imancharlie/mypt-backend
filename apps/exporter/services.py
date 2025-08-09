import os
import logging
from io import BytesIO
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.pdfbase.pdfmetrics import stringWidth
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.conf import settings
from django.template.loader import render_to_string
from django.http import HttpResponse

logger = logging.getLogger(__name__)

def export_weekly_report_pdf(weekly_report):
    """Export weekly report as PDF in CoET format."""
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from io import BytesIO
    
    buffer = BytesIO()
    # Match typical DOCX defaults: 1 inch margins
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )
    elements = []
    
    # Get styles
    styles = getSampleStyleSheet()
    
    # Create custom styles for professional format
    title_style = ParagraphStyle(
        'CoETTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=20,
        alignment=1,  # Center alignment
        fontName='Times-Bold',
        leading=19  # 1.5 line spacing
    )
    
    header_style = ParagraphStyle(
        'CoETHeader',
        parent=styles['Heading2'],
        fontSize=12,
        spaceAfter=12,
        fontName='Times-Bold',
        leading=18  # 1.5 line spacing
    )
    
    # Create normal text style
    normal_style = ParagraphStyle(
        'Normal',
        parent=styles['Normal'],
        fontSize=12,
        fontName='Times-Roman',
        leading=18,  # 1.5 line spacing
        spaceAfter=6
    )
    
    # Title
    title = Paragraph("College of Engineering and Technology (CoET)", title_style)
    elements.append(title)
    
    # Weekly Report Header
    daily_reports = weekly_report.get_daily_reports()
    total_hours = sum(report.hours_spent for report in daily_reports)
    
    report_header = Paragraph(f"Weekly Report No: {weekly_report.week_number} from: {weekly_report.start_date.strftime('%d-%m-%Y')} to: {weekly_report.end_date.strftime('%d-%m-%Y')}", header_style)
    elements.append(report_header)
    
    # Daily Reports Table - No "Weekly Work Summary" heading
    if daily_reports.exists():
        # Build table rows with Paragraphs for proper wrapping
        header_cell_style = ParagraphStyle(
            'TableHeader', parent=styles['Normal'], fontSize=12, fontName='Times-Bold', leading=18
        )
        cell_style = ParagraphStyle(
            'TableCell', parent=styles['Normal'], fontSize=12, fontName='Times-Roman', leading=18
        )

        # Day mapping
        day_names = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday']

        # Compute dynamic column widths so Day/Hours are tight and Description takes the rest
        available_width = doc.width  # points

        day_texts = ['Day']
        hours_texts = ['Hours', str(total_hours)]
        for i, report in enumerate(daily_reports):
            day_name = day_names[i] if i < len(day_names) else report.date.strftime('%A')
            day_texts.append(day_name)
            hours_texts.append(str(report.hours_spent))

        # Helper to compute minimal col width in points (content width + padding)
        def min_col_width(texts, font_name='Times-Roman', font_size=12, side_padding_pts=12):
            if not texts:
                return 0
            widest = max(stringWidth(str(t), font_name, font_size) for t in texts)
            return widest + side_padding_pts

        # Tight but readable columns; clamp to sensible bounds
        day_col = min(max(min_col_width(day_texts, 'Times-Roman', 12), 0.7*inch), 1.5*inch)
        hours_col = min(max(min_col_width(hours_texts, 'Times-Roman', 12), 0.6*inch), 1.2*inch)
        description_col = max(available_width - day_col - hours_col, 2.5*inch)

        data = [
            [
                Paragraph('Day', header_cell_style),
                Paragraph('Brief description of work performed', header_cell_style),
                Paragraph('Hours', header_cell_style),
            ]
        ]

        for i, report in enumerate(daily_reports):
            day_name = day_names[i] if i < len(day_names) else report.date.strftime('%A')
            data.append([
                Paragraph(day_name, cell_style),
                Paragraph(report.description or '', cell_style),
                Paragraph(str(report.hours_spent), cell_style),
            ])

        # Total row
        data.append([
            Paragraph('', cell_style),
            Paragraph('Total hrs', header_cell_style),
            Paragraph(str(total_hours), header_cell_style),
        ])

        daily_table = Table(data, colWidths=[day_col, description_col, hours_col], repeatRows=1)
        daily_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('FONTSIZE', (0, 0), (-1, -1), 12),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('REPEATROWS', (0, 0), (-1, 0)),
        ]))
        elements.append(daily_table)
        elements.append(Spacer(1, 20))
    
    # Main Job and Operations - Independent section
    if hasattr(weekly_report, 'main_job') and weekly_report.main_job:
        # Add space before main job section
        elements.append(Spacer(1, 20))
        
        # Create main job title as a paragraph (not table)
        main_job_title = weekly_report.main_job.title or "Sequence of Operations"
        main_job_title_style = ParagraphStyle(
            'MainJobTitle',
            parent=styles['Heading2'],
            fontSize=12,
            spaceAfter=12,
            fontName='Times-Bold',
            leading=18
        )
        main_job_title_para = Paragraph(main_job_title, main_job_title_style)
        elements.append(main_job_title_para)
        
        # Create operations table
        operations_data = [['No', 'Operation', 'Tools, Machinery, Equipment']]
        
        # Get actual operations from database - no mock data
        actual_operations = weekly_report.main_job.operations.all().order_by('step_number')
        
        if actual_operations.exists():
            # Use actual operations from database
            for operation in actual_operations:
                operations_data.append([
                    str(operation.step_number),
                    operation.operation_description or '',
                    operation.tools_used or ''
                ])
        else:
            # If no operations exist, show empty table with just header
            operations_data.append(['', '', ''])
        
        # Dynamically size columns so the Operation column gets the majority of width
        available_width = doc.width

        numbers = ['No'] + [str(op.step_number) for op in actual_operations]
        tools_samples = ['Tools, Machinery, Equipment'] + [
            (op.tools_used or '') for op in actual_operations
        ]

        no_col = min(max(min_col_width(numbers, 'Times-Roman', 12), 0.6*inch), 1.0*inch)
        tools_col = min(max(min_col_width(tools_samples, 'Times-Roman', 12, side_padding_pts=12), 1.0*inch), 1.8*inch)
        operation_col = max(available_width - no_col - tools_col, 3.0*inch)

        # Convert operations_data text to Paragraphs
        op_header = [
            Paragraph('No', header_cell_style),
            Paragraph('Operation', header_cell_style),
            Paragraph('Tools, Machinery, Equipment', header_cell_style),
        ]
        data_ops = [op_header]

        if actual_operations.exists():
            for op in actual_operations:
                data_ops.append([
                    Paragraph(str(op.step_number), cell_style),
                    Paragraph(op.operation_description or '', cell_style),
                    Paragraph(op.tools_used or '', cell_style),
                ])
        else:
            data_ops.append([
                Paragraph('', cell_style),
                Paragraph('', cell_style),
                Paragraph('', cell_style),
            ])

        operations_table = Table(data_ops, colWidths=[no_col, operation_col, tools_col], repeatRows=1)
        operations_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('FONTSIZE', (0, 0), (-1, -1), 12),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('REPEATROWS', (0, 0), (-1, 0)),
        ]))
        elements.append(operations_table)
        elements.append(Spacer(1, 30))
    
    # Signature section - Create a small table for signature (right half)
    signature_data = [
        ['', ''],
        ['Signature Training Officer', 'Date']
    ]
    
    # Create a container table to position signature on the right
    container_data = [['', '']]  # Empty first row for spacing
    container_table = Table(container_data, colWidths=[3.5*inch, 3.5*inch])
    container_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ('GRID', (0, 0), (-1, -1), 0, colors.white),  # No visible grid
    ]))
    
    # Create signature table (smaller size)
    signature_table = Table(signature_data, colWidths=[1.5*inch, 1.5*inch])
    signature_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Times-Roman'),
        ('FONTSIZE', (0, 0), (-1, -1), 12),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('WORDWRAP', (0, 0), (-1, -1), True),
        ('ROWBACKGROUNDS', (0, 0), (-1, -1), [colors.white, colors.white])
    ]))
    
    elements.append(Spacer(1, 20))
    elements.append(container_table)
    
    # Position signature table on the right side
    from reportlab.platypus import KeepTogether
    signature_container = KeepTogether([signature_table])
    elements.append(signature_container)
    
    # Build the PDF
    doc.build(elements)
    pdf = buffer.getvalue()
    buffer.close()
    return pdf

def export_weekly_report_docx(weekly_report):
    """Export weekly report as DOCX in CoET format."""
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from io import BytesIO
    
    doc = Document()
    
    # Set document margins to 1 inch to match PDF
    from docx.shared import Pt
    from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('College of Engineering and Technology (CoET)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Set font for title
    for run in title.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True
    
    # Set line spacing for title
    title.paragraph_format.line_spacing = 1.5
    
    # Weekly Report Header
    daily_reports = weekly_report.get_daily_reports()
    total_hours = sum(report.hours_spent for report in daily_reports)
    
    report_header = doc.add_heading(f'Weekly Report No: {weekly_report.week_number} from: {weekly_report.start_date.strftime("%d-%m-%Y")} to: {weekly_report.end_date.strftime("%d-%m-%Y")}', level=1)
    
    # Set font for header
    for run in report_header.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
    
    # Set line spacing for header
    report_header.paragraph_format.line_spacing = 1.5
    
    # Daily Reports Section - No "Weekly Work Summary" heading
    if daily_reports.exists():
        # Create daily reports table with dynamically computed column widths
        daily_table = doc.add_table(rows=1, cols=3)
        daily_table.style = 'Table Grid'

        # Compute available content width in inches
        section = doc.sections[0]
        # EMU per inch constant used by python-docx
        EMU_PER_INCH = 914400
        content_width_inches = (section.page_width - section.left_margin - section.right_margin) / EMU_PER_INCH

        # Build datasets to estimate tight widths
        day_names = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday']
        days = ['Day'] + [
            (day_names[i] if i < len(day_names) else r.date.strftime('%A'))
            for i, r in enumerate(daily_reports)
        ]
        hours_values = ['Hours'] + [str(r.hours_spent) for r in daily_reports] + ['00']

        # Approximate text width in inches for Times New Roman 12pt
        def approx_inch_width(text: str, per_char=0.09, base_padding=0.3) -> float:
            return max(base_padding + len(str(text)) * per_char, 0.0)

        day_col_in = min(max(max(approx_inch_width(t) for t in days), 0.7), 1.3)
        hours_col_in = min(max(max(approx_inch_width(t) for t in hours_values), 0.6), 1.1)
        description_col_in = max(content_width_inches - day_col_in - hours_col_in, 2.5)

        daily_table.columns[0].width = Inches(day_col_in)
        daily_table.columns[1].width = Inches(description_col_in)
        daily_table.columns[2].width = Inches(hours_col_in)
        
        # Header
        header_row = daily_table.rows[0]
        header_row.cells[0].text = 'Day'
        header_row.cells[1].text = 'Brief description of work performed'
        header_row.cells[2].text = 'Hours'
        
        # Day mapping
        day_names = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday']
        
        # Data rows
        for i, report in enumerate(daily_reports):
            row = daily_table.add_row()
            day_name = day_names[i] if i < len(day_names) else report.date.strftime('%A')
            row.cells[0].text = day_name
            row.cells[1].text = report.description
            row.cells[2].text = str(report.hours_spent)
            
            # Set font and formatting for table cells
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                    paragraph.alignment = WD_TABLE_ALIGNMENT.LEFT
            
            # Set word wrapping for description column to prevent overflow
            row.cells[1].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Add total row
        total_row = daily_table.add_row()
        total_row.cells[0].text = ''
        total_row.cells[1].text = 'Total hrs'
        total_row.cells[2].text = str(total_hours)
    
    # Main Job and Operations Section - Independent section
    if hasattr(weekly_report, 'main_job') and weekly_report.main_job:
        # Add space before main job section
        doc.add_paragraph()
        
        # Create main job title as a heading (not table)
        main_job_title = weekly_report.main_job.title or "Sequence of Operations"
        main_job_heading = doc.add_heading(main_job_title, level=1)
        
        # Set font for main job heading
        for run in main_job_heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
        
        # Set line spacing for main job heading
        main_job_heading.paragraph_format.line_spacing = 1.5
        
        # Create operations table with proper column widths
        operations_table = doc.add_table(rows=1, cols=3)
        operations_table.style = 'Table Grid'

        # Compute available width again
        content_width_inches = (section.page_width - section.left_margin - section.right_margin) / EMU_PER_INCH

        actual_operations = weekly_report.main_job.operations.all().order_by('step_number')
        numbers = ['No'] + [str(op.step_number) for op in actual_operations]
        tools_samples = ['Tools, Machinery, Equipment'] + [(op.tools_used or '') for op in actual_operations]

        no_col_in = min(max(max(approx_inch_width(t) for t in numbers), 0.6), 1.0)
        tools_col_in = min(max(max(approx_inch_width(t) for t in tools_samples), 1.0), 2.0)
        operation_col_in = max(content_width_inches - no_col_in - tools_col_in, 3.0)

        operations_table.columns[0].width = Inches(no_col_in)
        operations_table.columns[1].width = Inches(operation_col_in)
        operations_table.columns[2].width = Inches(tools_col_in)
        
        # Header
        header_row = operations_table.rows[0]
        header_row.cells[0].text = 'No'
        header_row.cells[1].text = 'Operation'
        header_row.cells[2].text = 'Tools, Machinery, Equipment'
        
        # Get actual operations from database - no mock data
        actual_operations = weekly_report.main_job.operations.all().order_by('step_number')
        
        if actual_operations.exists():
            # Use actual operations from database
            for operation in actual_operations:
                row = operations_table.add_row()
                row.cells[0].text = str(operation.step_number)
                row.cells[1].text = operation.operation_description or ''
                row.cells[2].text = operation.tools_used or ''
                
                # Set word wrapping for operation and tools columns
                row.cells[1].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.LEFT
                row.cells[2].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.LEFT
        else:
            # If no operations exist, show empty table with just header
            row = operations_table.add_row()
            row.cells[0].text = ''
            row.cells[1].text = ''
            row.cells[2].text = ''
    
    # Signature section - Create a small table for signature (right half)
    doc.add_paragraph()  # Add space before signature table
    
    # Create signature table (smaller size)
    signature_table = doc.add_table(rows=2, cols=2)
    signature_table.style = 'Table Grid'
    
    # Set column widths for signature table (smaller size)
    signature_table.columns[0].width = Inches(1.5)
    signature_table.columns[1].width = Inches(1.5)
    
    # Add signature content
    signature_table.rows[0].cells[0].text = ''
    signature_table.rows[0].cells[1].text = ''
    signature_table.rows[1].cells[0].text = 'Signature Training Officer'
    signature_table.rows[1].cells[1].text = 'Date'
    
    # Set font for signature table
    for row in signature_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                paragraph.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    docx_content = buffer.getvalue()
    buffer.close()
    return docx_content

def export_general_report_pdf(general_report):
    """Export general report to PDF format."""
    try:
        context = {
            'general_report': general_report,
        }
        
        html_content = render_to_string('exports/general_report.html', context)
        
        # Create a simple text response for now
        response = HttpResponse(html_content, content_type='text/html')
        response['Content-Disposition'] = f'attachment; filename="general_report.html"'
        return response
        
    except Exception as e:
        logger.error(f"Error exporting general report PDF: {e}")
        raise

def export_general_report_docx(general_report):
    """Export general report to DOCX format."""
    try:
        doc = Document()
        
        # Add title
        title = doc.add_heading(general_report.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add sections
        sections = [
            ('Introduction', general_report.introduction),
            ('Company Overview', general_report.company_overview),
            ('Training Objectives', general_report.training_objectives),
            ('Methodology', general_report.methodology),
            ('Achievements', general_report.achievements),
            ('Challenges Faced', general_report.challenges_faced),
            ('Skills Acquired', general_report.skills_acquired),
            ('Recommendations', general_report.recommendations),
            ('Conclusion', general_report.conclusion),
            ('Acknowledgments', general_report.acknowledgments),
        ]
        
        for section_title, section_content in sections:
            if section_content:
                doc.add_heading(section_title, level=1)
                doc.add_paragraph(section_content)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        docx_content = buffer.getvalue()
        buffer.close()
        return docx_content
        
    except Exception as e:
        logger.error(f"Error exporting general report DOCX: {e}")
        raise 